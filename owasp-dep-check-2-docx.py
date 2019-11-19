import csv
import logging
import os
from pathlib import Path

from docx import Document

LOGGING_FORMAT = '%(asctime)s | %(levelname)8s | %(message)s'
LOGGING_TIME_AND_DATE_FORMAT = '%d/%b/%Y %a %H:%M:%S %z'

OWASP_DEP_CHECK_REPORTS_EXTENSION = '.csv'
INPUT_DIRECTORY_WITH_REPORTS = './owasp'
OUTPUT_FILE_EXTENSION = '.docx'
OUTPUT_DIRECTORY = 'output'


# More checks can be implemented
def is_suitable_owasp_report_file(file_to_check: str) -> bool:
    extension = os.path.splitext(file_to_check)[-1].lower()
    return extension == OWASP_DEP_CHECK_REPORTS_EXTENSION


def get_severity(csv_row):
    if csv_row['CVSSv3_BaseSeverity'] == "" or csv_row['CVSSv3_BaseSeverity'] is None:
        logging.warning("Found vulnerability with no CVSSv3 severity score, using v2 instead")
        return csv_row['CVSSv2_Severity']
    else:
        return csv_row['CVSSv3_BaseSeverity']


def process_owasp_report_to_docx(owasp_report_path: str):
    logging.info(f'Processing OWASP Dependency Check report file to docx: {owasp_report_path}')

    logging.info('Creating Word document object')
    document = Document()
    document.add_heading(owasp_report_path.upper())
    document.add_page_break()

    logging.info("Preparing docx table with vulnerabilities")
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Vulnerable Component'
    hdr_cells[1].text = 'Severity'
    hdr_cells[2].text = 'CVE Identifier'
    hdr_cells[3].text = 'Vulnerability Description'

    logging.info(f'Opening OWASP report file to parse vulnerabilities: {owasp_report_path}')
    with open(owasp_report_path) as report_file:
        logging.info("Reading vulnerabilities information from OWASP report file")
        csv_report_reader = csv.DictReader(report_file)
        for csv_row in csv_report_reader:
            row_cells = table.add_row().cells
            row_cells[0].text = csv_row['DependencyName']  # Component
            row_cells[1].text = get_severity(csv_row)  # Severity
            row_cells[2].text = csv_row['CVE']  # CVE
            row_cells[3].text = csv_row['Vulnerability']  # Description

    output_file = os.path.join(OUTPUT_DIRECTORY, Path(file).with_suffix(OUTPUT_FILE_EXTENSION))
    logging.info(f'Finished processing. Saving to: {output_file}')
    document.save(output_file)


if __name__ == '__main__':
    logging.basicConfig(level=logging.INFO, format=LOGGING_FORMAT, datefmt=LOGGING_TIME_AND_DATE_FORMAT)
    logging.info(f'Creating output directory: {OUTPUT_DIRECTORY}')
    os.makedirs(OUTPUT_DIRECTORY, exist_ok=True)

    # Traverse the input directory recursively
    for root_directory, subdirectories, files_list in os.walk(INPUT_DIRECTORY_WITH_REPORTS):
        for file in files_list:
            file_abs_path = os.path.join(os.path.abspath(root_directory), file)
            if is_suitable_owasp_report_file(file):
                logging.info(f'Found OWASP report file: {file_abs_path}')
                process_owasp_report_to_docx(file_abs_path)
            else:
                logging.warning(f'Skipping file: {file_abs_path}')
